import os
import csv
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import List, Tuple
import sqlite3
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF, XPos, YPos
import zipfile
import rarfile
import py7zr
import PyPDF2

def generate_sample_files(root_dir: str = "document_storage"):
    """
    Создаёт структуру папок и набор файлов разных форматов
    с осмысленным текстом для демонстрации поиска.
    """
    os.makedirs(root_dir, exist_ok=True)
    subdirs = ["docs", "spreadsheets", "pdfs", "archives"]
    for sub in subdirs:
        os.makedirs(os.path.join(root_dir, sub), exist_ok=True)

    doc = Document()
    doc.add_heading("Финансовый отчёт компании", level=1)
    doc.add_paragraph("Этот документ содержит данные о прибыли за 1 квартал 2025 года.")
    doc.add_paragraph("Выручка составила 10 миллионов рублей, чистая прибыль — 2.5 миллиона.")
    doc.save(os.path.join(root_dir, "docs", "financial_report.docx"))

    doc2 = Document()
    doc2.add_heading("Техническая документация", level=1)
    doc2.add_paragraph("API позволяет интегрировать поисковый движок с внешними сервисами.")
    doc2.add_paragraph("Используйте endpoint /search для выполнения запросов.")
    doc2.save(os.path.join(root_dir, "docs", "api_manual.docx"))

    wb = Workbook()
    ws = wb.active
    ws.title = "Продажи"
    ws.append(["Товар", "Количество", "Цена", "Сумма"])
    ws.append(["Ноутбук", 5, 70000, 350000])
    ws.append(["Монитор", 10, 15000, 150000])
    ws.append(["Клавиатура", 20, 2500, 50000])
    ws.append(["Итого: общая выручка 550000 рублей."])
    wb.save(os.path.join(root_dir, "spreadsheets", "sales_data.xlsx"))

    wb2 = Workbook()
    ws2 = wb2.active
    ws2.title = "Сотрудники"
    ws2.append(["ID", "Имя", "Должность"])
    ws2.append([1, "Иван Иванов", "Менеджер"])
    ws2.append([2, "Петр Петров", "Разработчик"])
    wb2.save(os.path.join(root_dir, "spreadsheets", "employees.xlsx"))

    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("DejaVu", "", "fonts/DejaVuSans.ttf")
    pdf.set_font("DejaVu", size=12)
    pdf.cell(200, 10, text="Маркетинговый план на 2025 год", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(10)
    pdf.multi_cell(0, 10, text="Основные цели: увеличить узнаваемость бренда, запустить рекламную кампанию в социальных сетях, провести ребрендинг.")
    pdf.output(os.path.join(root_dir, "pdfs", "marketing_plan.pdf"))

    pdf2 = FPDF()
    pdf2.add_page()
    pdf2.add_font("DejaVu", "", "fonts/DejaVuSans.ttf")
    pdf2.set_font("DejaVu", size=12)
    pdf2.cell(200, 10, text="Инструкция по безопасности", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf2.multi_cell(0, 10, text="Все сотрудники должны использовать надёжные пароли. Доступ к серверам только через VPN.")
    pdf2.output(os.path.join(root_dir, "pdfs", "security_policy.pdf"))

    with open(os.path.join(root_dir, "docs", "notes.txt"), "w", encoding="utf-8") as f:
        f.write("Заметки со встречи:\n")
        f.write("- Обсудили бюджет на разработку новой поисковой системы.\n")
        f.write("- Срок сдачи прототипа — 1 июля 2025.\n")

    archive_files = [
        os.path.join(root_dir, "docs", "financial_report.docx"),
        os.path.join(root_dir, "pdfs", "security_policy.pdf"),
        os.path.join(root_dir, "spreadsheets", "sales_data.xlsx")
    ]

    with zipfile.ZipFile(os.path.join(root_dir, "archives", "bundle.zip"), "w") as zf:
        for f in archive_files:
            zf.write(f, os.path.basename(f))

    rar_path = shutil.which("rar") or shutil.which("winrar")
    if rar_path:
        rar_path = str(rar_path)
        rar_exe = rar_path if rar_path.lower().endswith("rar.exe") else os.path.join(os.path.dirname(rar_path),
                                                                                     "Rar.exe")
        if os.path.exists(rar_exe):
            subprocess.run([rar_exe, "a", os.path.join(root_dir, "archives", "bundle.rar")] + archive_files, check=True,
                           cwd=root_dir)
        else:
            print("RAR не найден, создание архива .rar пропущено.")
    else:
        print("RAR не найден, создание архива .rar пропущено.")

    with py7zr.SevenZipFile(os.path.join(root_dir, "archives", "bundle.7z"), "w") as sz:
        for f in archive_files:
            sz.write(f, os.path.basename(f))

    print(f"Тестовые файлы созданы в '{root_dir}'")

def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_from_xlsx(file_path: str) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(file_path, data_only=True)
    text_parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text_parts.append(f"[Лист: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            row_text = " ".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip():
                text_parts.append(row_text)
    return "\n".join(text_parts)

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="cp1251") as f:
            return f.read()

def extract_text_from_file(file_path: str) -> str:
    """Определяет тип файла по расширению и извлекает текст."""
    ext = Path(file_path).suffix.lower()
    try:
        if ext == ".docx":
            return extract_text_from_docx(file_path)
        elif ext == ".xlsx":
            return extract_text_from_xlsx(file_path)
        elif ext == ".pdf":
            return extract_text_from_pdf(file_path)
        elif ext == ".txt":
            return extract_text_from_txt(file_path)
        elif ext in (".doc", ".xls"):
            # Старые форматы пропускаем
            return ""
        else:
            return ""
    except Exception as e:
        return f"ERROR extracting {file_path}: {e}"


def process_archive(archive_path: str, temp_dir: str) -> List[Tuple[str, str]]:
    """
    Извлекает файлы из архива и возвращает список кортежей:
      (логический_путь, извлечённый_текст),
    где логический_путь = 'путь_к_архиву/внутренний_путь_в_архиве'.
    """
    results_list = []
    ext = Path(archive_path).suffix.lower()
    extract_root = os.path.join(temp_dir, "archive_content")
    os.makedirs(extract_root, exist_ok=True)

    file_list = []
    try:
        if ext == ".zip":
            with zipfile.ZipFile(archive_path, "r") as zf:
                file_list = [name for name in zf.namelist() if not name.endswith("/")]
        elif ext == ".rar":
            with rarfile.RarFile(archive_path, "r") as rf:
                file_list = [info.filename for info in rf.infolist() if not info.isdir()]
        elif ext == ".7z":
            with py7zr.SevenZipFile(archive_path, "r") as sz:
                file_list = [name for name in sz.getnames() if not name.endswith("/")]
        else:
            return results_list
    except Exception as e:
        print(f"Ошибка чтения архива {archive_path}: {e}")
        return results_list

    try:
        if ext == ".zip":
            with zipfile.ZipFile(archive_path, "r") as zf:
                zf.extractall(extract_root)
        elif ext == ".rar":
            with rarfile.RarFile(archive_path, "r") as rf:
                rf.extractall(extract_root)
        elif ext == ".7z":
            with py7zr.SevenZipFile(archive_path, "r") as sz:
                sz.extractall(extract_root)
    except Exception as e:
        print(f"Ошибка распаковки {archive_path}: {e}")
        return results_list

    for internal_name in file_list:
        extracted_file_path = os.path.normpath(os.path.join(extract_root, internal_name))
        if os.path.isfile(extracted_file_path):
            content = extract_text_from_file(extracted_file_path)
            if content.strip():
                logical_path = archive_path.replace("\\", "/") + "/" + internal_name.replace("\\", "/")
                results_list.append((logical_path, content))
        else:
            print(f"После распаковки не найден файл: {internal_name} в архиве {archive_path}")

    return results_list


def crawl_storage(root_dir: str, output_csv: str = "parsed_documents.csv"):
    """
    Рекурсивно обходит root_dir, извлекает текст из поддерживаемых файлов
    и сохраняет результат в CSV.
    """
    with open(output_csv, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(["file_path", "file_type", "content"])

        with tempfile.TemporaryDirectory() as temp_dir:
            for current_root, dirs, files in os.walk(root_dir):
                for file in files:
                    full_path = os.path.join(current_root, file)
                    file_type = Path(file).suffix.lower()
                    if file.startswith("~$"):
                        continue
                    if file_type in (".zip", ".rar", ".7z"):
                        archived_files = process_archive(full_path, temp_dir)
                        for arc_path, arc_content in archived_files:
                            writer.writerow([arc_path, Path(arc_path).suffix, arc_content])
                    else:
                        content = extract_text_from_file(full_path)
                        if content.strip():
                            writer.writerow([full_path, file_type, content])
    print(f"Результаты краулера сохранены в '{output_csv}'")


def create_and_populate_fts_db(csv_file: str, db_file: str = "documents_fts.db"):
    """Создаёт SQLite БД с виртуальной таблицей FTS5 и загружает данные из CSV."""
    conn = sqlite3.connect(db_file)
    cursor = conn.cursor()

    cursor.execute("""
        CREATE VIRTUAL TABLE IF NOT EXISTS docs_fts USING fts5(
            file_path,
            file_type,
            content,
            tokenize='unicode61'
        )
    """)

    with open(csv_file, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            cursor.execute(
                "INSERT INTO docs_fts (file_path, file_type, content) VALUES (?, ?, ?)",
                (row["file_path"], row["file_type"], row["content"])
            )

    conn.commit()
    conn.close()
    print(f"База данных '{db_file}' создана и наполнена.")


def search_fts(query: str, db_file: str = "documents_fts.db") -> list:
    """Выполняет полнотекстовый поиск и возвращает список найденных записей."""
    conn = sqlite3.connect(db_file)
    cursor = conn.cursor()
    cursor.execute("""
        SELECT file_path, snippet(docs_fts, 2, '<b>', '</b>', '...', 40)
        FROM docs_fts
        WHERE docs_fts MATCH ?
    """, (query,))
    results_list = cursor.fetchall()
    conn.close()
    return results_list


if __name__ == "__main__":
    STORAGE_DIR = "document_storage"
    CSV_OUTPUT = "parsed_documents.csv"
    DB_FILE = "documents_fts.db"

    if not os.path.exists(STORAGE_DIR) or not os.listdir(STORAGE_DIR):
        generate_sample_files(STORAGE_DIR)
    else:
        print(f"Хранилище '{STORAGE_DIR}' уже существует, генерация пропущена.")

    crawl_storage(STORAGE_DIR, CSV_OUTPUT)

    create_and_populate_fts_db(CSV_OUTPUT, DB_FILE)

    print("\nПример поиска по слову 'сотрудники':")
    results = search_fts("сотрудники")
    for path, snippet in results:
        print(f" {path}\n     ...{snippet}...\n")

    print("\nПример поиска по слову 'прибыль':")
    results = search_fts("прибыль")
    for path, snippet in results:
        print(f" {path}\n     ...{snippet}...\n")